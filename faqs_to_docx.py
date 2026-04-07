#!/usr/bin/env python3
"""
Genera `FAQS.docx` a partir de `data/faqs_normalized.json`.
"""
import json
from pathlib import Path
from docx import Document
from docx.shared import Pt


def load_faqs(path=None):
    # prefer data/faqs_final.json > data/faqs_normalized.json > data/faqs.json
    candidates = []
    if path:
        candidates.append(Path(path))
    candidates.extend([Path('data/faqs_final.json'), Path('data/faqs_normalized.json'), Path('data/faqs.json')])
    for p in candidates:
        if p and p.exists():
            return json.loads(p.read_text(encoding='utf-8'))
    raise FileNotFoundError('No se encontró ningún fichero de FAQs (buscando data/faqs_final.json, data/faqs_normalized.json, data/faqs.json)')


def write_docx(faqs, out_path='FAQS.docx'):
    doc = Document()
    doc.core_properties.title = 'FAQs extraídas del SII'
    doc.add_heading('Preguntas Frecuentes (extraídas)', level=1)

    for i, f in enumerate(faqs, start=1):
        q = f.get('normalized_question') or f.get('question')
        a = f.get('answer') or ''
        src = f.get('source','')
        url = f.get('path','')

        doc.add_heading(f'{i}. {q}', level=2)
        p_meta = doc.add_paragraph()
        p_meta.add_run('Fuente: ').bold = True
        p_meta.add_run(str(src) + '\n')
        p_meta.add_run('Origen (archivo): ').bold = True
        p_meta.add_run(str(url) + '\n')

        # respuesta
        p = doc.add_paragraph()
        p.add_run('Respuesta:').bold = True
        p = doc.add_paragraph(a)
        # spacing
        doc.add_paragraph('')

    # aplicar estilo base (opcional)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    doc.save(out_path)
    print(f'Documento generado: {out_path}')


def main():
    faqs = load_faqs()
    write_docx(faqs)


if __name__ == '__main__':
    main()
